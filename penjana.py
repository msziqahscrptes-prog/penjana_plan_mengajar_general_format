import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO

# --- 1. CONFIGURATION ---
st.set_page_config(page_title="Penjana RPH Bahasa Melayu", layout="wide")
st.title("🎓 PENJANA RANCANGAN PENGAJARAN HARIAN (RPH)")

# --- MAIN PAGE CONFIGURATION & USER API KEY BAR (AT THE VERY TOP) ---
user_api_key = st.text_input(
    "🔑 MASUKKAN KUNCI API GEMINI ANDA:", 
    type="password", 
    help="Dapatkan kunci API anda dari Google AI Studio menggunakan akaun Gmail anda."
)

# Helper function to dynamically check and load models based on the user's key
def get_working_model(api_key):
    try:
        genai.configure(api_key=api_key)
        for m in genai.list_models():
            if 'generateContent' in m.supported_generation_methods:
                return m.name
    except Exception as e:
        st.error(f"KUNCI API TIDAK SAH ATAU RALAT SAMBUNGAN: {str(e)}")
        return None
    return "models/gemini-1.5-flash"  # Default fallback


# Process model assignment if the key is provided
selected_model_name = None
if user_api_key:
    selected_model_name = get_working_model(user_api_key)
    if selected_model_name:
        st.info(f"SISTEM DISAMBUNGKAN. MODEL AKTIF: {selected_model_name.upper()}")
else:
    st.warning("⚠️ SILA MASUKKAN KUNCI API GEMINI PERIBADI ANDA DI ATAS UNTUK BERMULA.")


# --- 2. AI GENERATION ENGINE (BAHASA MELAYU SPECIFIC) ---
def generate_bm_plan(topic, syllabus, extra_context, api_key, model_name):
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel(model_name)
    
    prompt = f"""
    Topik: {topic}. Kod Sukatan/Standard Kandungan: {syllabus}. Konteks Tambahan: {extra_context}.
    Hasilkan satu Rancangan Pengajaran Harian (RPH) yang profesional dalam Bahasa Melayu sepenuhnya.
    
    PERATURAN FORMAT TEKS KRITIKAL:
    1. JANGAN gunakan dua tanda asteris (**) di mana-mana bahagian output.
    2. Pastikan setiap tajuk bahagian (SECTION:) ditulis dalam HURUF BESAR SEPENUHNYA.
    3. JANGAN gunakan senarai bulet atau titik. Gunakan sistem penomboran (1, 2, 3...) untuk semua senarai item.
    
    Strukturkan output dengan penanda eksak berikut untuk pembetulan kotak dokumen:
    
    SECTION: TOPIK DAN KOD SUKATAN
    {topic.upper()} ({syllabus.upper()})
    
    SECTION: OBJEKTIF PEMBELAJARAN
    [4 mata menggunakan nombor 1 hingga 4]
    
    SECTION: HASIL PEMBELAJARAN
    [4 mata menggunakan nombor 1 hingga 4]
    
    SECTION: KRITERIA KEJAYAAN
    [4 mata menggunakan nombor 1 hingga 4]
    
    SECTION: PENGETAHUAN SEDIA ADA
    [1 mata menggunakan nombor 1]
    
    SECTION: KATA KUNCI
    [6 item menggunakan nombor 1 hingga 6]
    
    SECTION: SOALAN KEMAHIRAN BERFIKIR ARAS TINGGI (KBAT)
    [4 soalan KBAT yang berbeza menggunakan nombor 1 hingga 4]
    
    SECTION: STRATEGI PENGAJARAN TERBEZA (KUMPULAN HIJAU - TINGGI)
    1. Aktiviti Terbeza Murid Pencapaian Tinggi: [Penerangan aktiviti]
    
    SECTION: STRATEGI PENGAJARAN TERBEZA (KUMPULAN KUNING - SEDERHANA)
    1. Aktiviti Terbeza Murid Pencapaian Sederhana: [Penerangan aktiviti]
    
    SECTION: STRATEGI PENGAJARAN TERBEZA (KUMPULAN MERAH - RENDAH)
    1. Aktiviti Terbeza Murid Pencapaian Rendah: [Penerangan aktiviti]
    
    SECTION: PENTAKSIRAN MURID PENCAPAIAN TINGGI
    [1 bentuk pentaksiran/tugasan khusus menggunakan nombor 1]
    
    SECTION: PENTAKSIRAN MURID PENCAPAIAN SEDERHANA
    [1 bentuk pentaksiran/tugasan khusus menggunakan nombor 1]
    
    SECTION: PENTAKSIRAN MURID PENCAPAIAN RENDAH
    [1 bentuk pentaksiran/tugasan khusus menggunakan nombor 1]
    
    SECTION: AKTIVITI PEMBELAJARAN TERADUN SATU (15 MINIT)
    [Penerangan aktiviti berasaskan teknologi/digital, persediaan guru, dan tugasan murid menggunakan nombor 1, 2, 3...]
    
    SECTION: AKTIVITI PEMBELAJARAN TERADUN DUA (15 MINIT)
    [Penerangan aktiviti berasaskan teknologi/digital, persediaan guru, dan tugasan murid menggunakan nombor 1, 2, 3...]
    
    SECTION: PLENARI
    [Aktiviti penutup/rumusan pengajaran menggunakan nombor 1, 2...]
    
    SECTION: TUGASAN RUMAH
    [Tugasan susulan untuk murid menggunakan nombor 1]
    """
    try:
        response = model.generate_content(prompt)
        return response.text.replace("**", "")
    except Exception as e:
        return f"RALAT SISTEM: {str(e)}"


# --- 3. WORD DOCUMENT EXPORT ENGINE (CUSTOM PAGE SIZE, MARGINS, AND HEADERS) ---
def add_page_number(run):
    """Helper function to inject dynamic page numbering XML into Word header"""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def create_word_export(topic, syllabus, text):
    doc = Document()
    
    # 3a. Page Setup (Portrait, Width: 8.5", Height: 11.5", Margins: 0.4")
    section_geo = doc.sections[0]
    section_geo.page_width = Inches(8.5)
    section_geo.page_height = Inches(11.5)
    section_geo.top_margin = Inches(0.4)
    section_geo.bottom_margin = Inches(0.4)
    section_geo.left_margin = Inches(0.4)
    section_geo.right_margin = Inches(0.4)
    
    # 3b. Dynamic Page Numbering at TOP and MIDDLE alignment
    header = section_geo.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_p.add_run()
    header_run.font.name = 'Arial'
    header_run.font.size = Pt(10)
    add_page_number(header_run)

    # 3c. Global Typography Styles Override (Arial, 14pt, 1.0 Line Spacing)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(14)  # 14 points font size enforced
    
    p_format = style.paragraph_format
    p_format.line_spacing = 1.0  # 1 paragraph / single spacing enforced
    p_format.space_after = Pt(12)

    # Document Title Block
    title_p = doc.add_paragraph()
    run_title = title_p.add_run(f'RANCANGAN PENGAJARAN HARIAN: {topic.upper()}')
    run_title.bold = True
    run_title.font.size = Pt(16)

    # 3d. Administrative Header Table
    admin_table = doc.add_table(rows=3, cols=4)
    admin_table.style = 'Table Grid'
    labels = [
        ["TARIKH:", "HARI:"], 
        ["TEMPAT / BILIK:", "MINGGU NO:"], 
        ["BILANGAN MURID:", "TEMPOH (MINIT):"]
    ]
    for r in range(3):
        admin_table.cell(r, 0).paragraphs[0].add_run(labels[r][0]).bold = True
        admin_table.cell(r, 2).paragraphs[0].add_run(labels[r][1]).bold = True
    doc.add_paragraph()

    # 3e. Content Parsing & Table Boxing
    sections = text.split('SECTION:')
    for section in sections:
        if not section.strip(): 
            continue
            
        lines = section.strip().split('\n')
        title = lines[0].strip().upper().replace("**", "")  # Enforce FULL CAPITAL LETTERS
        body_content = "\n".join(lines[1:])

        p_sec = doc.add_paragraph()
        p_sec.add_run(title).bold = True

        table = doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        cell_p = table.cell(0, 0).paragraphs[0]
        cell_p.paragraph_format.line_spacing = 1.0
        
        cleaned_body = body_content.strip().replace("**", "")
        cell_p.add_run(cleaned_body if cleaned_body else "Tiada kandungan dihasilkan.")
        doc.add_paragraph()

    # 3f. Principal / HOD Verification Block at the end
    doc.add_page_break()
    p_hod = doc.add_paragraph()
    p_hod.add_run("PENGESAHAN DAN ULASAN PENGETUA / KETUA JABATAN (HOD)").bold = True
    
    hod_table = doc.add_table(rows=3, cols=2)
    hod_table.style = 'Table Grid'
    hod_table.cell(0, 0).paragraphs[0].add_run("ULASAN / REMARK:").bold = True
    hod_table.cell(0, 1).paragraphs[0].add_run("TANDATANGAN & COP JABATAN:").bold = True
    hod_table.rows[1].height = Pt(50)  # Empty blank height space for stamping/signing
    hod_table.cell(2, 0).paragraphs[0].add_run("TARIKH SEMAKAN:").bold = True
    hod_table.cell(2, 1).paragraphs[0].add_run("NAMA PENYEMAK:").bold = True

    # Adjust inner spacing across cell tables formatting parameters safely
    for row in admin_table.rows:
        for cell in row.cells: cell.paragraphs[0].paragraph_format.line_spacing = 1.0
    for row in hod_table.rows:
        for cell in row.cells: cell.paragraphs[0].paragraph_format.line_spacing = 1.0

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


# --- 4. MAIN USER INTERFACE GRAPHICS ---
st.write("---")
u_topic = st.text_input("TOPIK / TAJUK PELAJARAN:")
u_syllabus = st.text_input("KOD SUKATAN / STANDARD KANDUNGAN:")
u_extra = st.text_area("KONTEKS TAMBAHAN / BAHAN BANTU BELAJAR (PILIHAN):")

if st.button("🚀 JANA RANCANGAN PENGAJARAN HARIAN", type="primary"):
    if not user_api_key:
        st.error("❌ RALAT KONFIGURASI! SILA MASUKKAN KUNCI API GEMINI ANDA DI BAHAGIAN ATAS HALAMAN TERLEBIH DAHULU.")
    elif not u_topic or not u_syllabus:
        st.error("❌ SILA ISI KEDUA-DUA RUANGAN TOPIK DAN KOD SUKATAN PELAJARAN.")
    else:
        with st.spinner("AI Sedang merangka RPH Bahasa Melayu anda..."):
            result = generate_bm_plan(u_topic, u_syllabus, u_extra, user_api_key, selected_model_name)
            st.session_state['bm_lp_out'] = result

if 'bm_lp_out' in st.session_state:
    st.divider()
    st.subheader("👁️ PRATONTON AI (PREVIEW COLUMNS)")
    st.text_area("PREVIEW KANDUNGAN RPH", st.session_state['bm_lp_out'], height=350)
    
    doc_file = create_word_export(u_topic, u_syllabus, st.session_state['bm_lp_out'])
    st.download_button(
        label="📥 MUAT TURUN VERSI WORD RASMI (.DOCX)", 
        data=doc_file, 
        file_name=f"RPH_BM_{u_topic.upper().replace(' ', '_')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# --- FOOTER SECTION ---
st.markdown("---") 
st.markdown(
    """
    <div style='text-align: center; color: grey; font-size: 0.8em;'>
        <p><b>Sistem Penjana RPH Pintar Bahasa Melayu v1.0</b></p>
        <p>Dibangunkan & Dikonsepkan oleh: <b>[Hajah Nurul Haziqah @ Hjh Hartini Hj Nordin]</b></p>
        <p>© 2026 BSc(Honors) in Computer Science, University of Strathclyde</p>
    </div>
    """,
    unsafe_allow_html=True
)
